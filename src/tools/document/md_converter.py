import os
import re
from docx import Document


def map_path(path: str) -> str:
    if os.environ.get("AMEVA_IN_CONTAINER") == "true":
        normalized = path.replace("\\", "/")
        return re.sub(r'^[Cc]:/ameva', '/app/workspace', normalized)
    return path


def convert_md_to_docx_logic(input_md_path: str, output_docx_path: str) -> str:
    """
    마크다운을 DOCX로 변환합니다.
    MCP 의존성이 전혀 없는 순수 파이썬 함수 (느슨한 결합)
    헤딩, 불릿, 번호목록, 코드블록, 굵은글씨, 수평선 지원.
    """
    orig_input = input_md_path
    orig_output = output_docx_path
    input_md_path = map_path(input_md_path)
    output_docx_path = map_path(output_docx_path)
    
    out_dir = os.path.dirname(output_docx_path)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    if not os.path.exists(input_md_path):
        return f"Error: Input file does not exist at {orig_input} (mapped to {input_md_path})"

    try:
        doc = Document()
        in_code_block = False
        code_lines = []
        numbered_counter = 0

        with open(input_md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        for raw_line in lines:
            line = raw_line.rstrip()
            
            # 코드 블록 처리
            if line.startswith("```"):
                if in_code_block:
                    # 코드 블록 종료
                    code_text = "\n".join(code_lines)
                    p = doc.add_paragraph(style="No Spacing")
                    run = p.add_run(code_text)
                    run.font.name = "Courier New"
                    run.font.size = __import__("docx.shared", fromlist=["Pt"]).Pt(9) if False else None
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue
            
            if in_code_block:
                code_lines.append(line)
                continue

            # 빈 줄
            if not line.strip():
                numbered_counter = 0
                continue

            # 헤딩
            if line.startswith("#### "):
                doc.add_heading(line[5:].strip(), level=4)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            # 수평선
            elif line.strip() in ["---", "***", "___"]:
                doc.add_paragraph("─" * 50)
            # 불릿 리스트
            elif line.startswith("- ") or line.startswith("* ") or line.startswith("+ "):
                text = line[2:].strip()
                # 볼드 처리
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                doc.add_paragraph(text, style='List Bullet')
            # 번호 리스트
            elif re.match(r'^\d+\. ', line):
                text = re.sub(r'^\d+\. ', '', line)
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                doc.add_paragraph(text, style='List Number')
            # 인용문
            elif line.startswith("> "):
                text = line[2:].strip()
                p = doc.add_paragraph(style="Quote" if "Quote" in [s.name for s in doc.styles] else "Normal")
                p.add_run(f'"{text}"').italic = True
            # 일반 텍스트 (볼드 처리 포함)
            else:
                p = doc.add_paragraph()
                # **bold** 파싱
                parts = re.split(r'\*\*(.+?)\*\*', line)
                for i, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.bold = (i % 2 == 1)
                
        doc.save(output_docx_path)
        return f"Success: Converted {orig_input} to {orig_output}"
        
    except Exception as e:
        return f"Error during conversion: {str(e)}"


def docx_to_markdown(docx_path: str, output_md_path: str = None) -> str:
    """
    .docx 파일을 마크다운(.md)으로 변환한다.
    헤딩 스타일, 리스트, 일반 단락을 파싱하여 구조화된 MD로 저장.
    output_md_path가 없으면 결과 텍스트를 직접 반환.
    """
    norm_path = map_path(docx_path)
    
    # 보안 검사
    abs_path = os.path.abspath(norm_path)
    if not abs_path.lower().startswith(r"c:\ameva") and \
       not abs_path.lower().startswith("/app/workspace"):
        return f"Security Error: Access to path '{abs_path}' is denied."
    
    if not os.path.exists(abs_path):
        return f"Error: DOCX file not found at {docx_path}"

    try:
        doc = Document(abs_path)
        md_lines = []

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else "Normal"
            text = para.text.strip()
            
            if not text:
                md_lines.append("")
                continue

            # 헤딩 스타일
            if "Heading 1" in style_name:
                md_lines.append(f"# {text}")
            elif "Heading 2" in style_name:
                md_lines.append(f"## {text}")
            elif "Heading 3" in style_name:
                md_lines.append(f"### {text}")
            elif "Heading 4" in style_name:
                md_lines.append(f"#### {text}")
            elif "Heading 5" in style_name or "Heading 6" in style_name:
                md_lines.append(f"##### {text}")
            # 리스트 스타일
            elif "List Bullet" in style_name:
                md_lines.append(f"- {text}")
            elif "List Number" in style_name:
                md_lines.append(f"1. {text}")
            # 코드 스타일
            elif "Code" in style_name or "No Spacing" in style_name:
                md_lines.append(f"```\n{text}\n```")
            # 인용
            elif "Quote" in style_name:
                md_lines.append(f"> {text}")
            else:
                # 볼드/이탤릭 처리
                md_text = ""
                for run in para.runs:
                    r_text = run.text
                    if run.bold and run.italic:
                        r_text = f"***{r_text}***"
                    elif run.bold:
                        r_text = f"**{r_text}**"
                    elif run.italic:
                        r_text = f"*{r_text}*"
                    md_text += r_text
                md_lines.append(md_text if md_text.strip() else text)

        # 표 처리
        for table in doc.tables:
            if not table.rows:
                continue
            header = [cell.text.strip() for cell in table.rows[0].cells]
            md_lines.append("\n| " + " | ".join(header) + " |")
            md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                md_lines.append("| " + " | ".join(cells) + " |")
            md_lines.append("")

        result = "\n".join(md_lines)
        # 연속 빈 줄 정리
        result = re.sub(r"\n{3,}", "\n\n", result).strip()

        if output_md_path:
            out_norm = map_path(output_md_path)
            out_abs = os.path.abspath(out_norm)
            if not out_abs.lower().startswith(r"c:\ameva") and \
               not out_abs.lower().startswith("/app/workspace"):
                return f"Security Error: Output path '{out_abs}' is denied."
            os.makedirs(os.path.dirname(out_abs), exist_ok=True) if os.path.dirname(out_abs) else None
            with open(out_abs, "w", encoding="utf-8") as f:
                f.write(result)
            return f"Success: Converted {docx_path} to {output_md_path} ({len(result)} chars)"
        
        # 직접 반환 (3000자 제한)
        preview = result[:3000]
        if len(result) > 3000:
            preview += f"\n\n... (truncated, full length: {len(result)} chars)"
        return preview

    except Exception as e:
        return f"Error converting DOCX to Markdown: {str(e)}"


def md_image_path_fixer(doc_path: str, base_image_dir: str) -> str:
    """
    마크다운 파일 내의 깨진 이미지 경로를 실제 로컬 이미지 경로로 자동 치환한다.
    base_image_dir 하위에서 동일한 파일명을 탐색하여 경로를 교정한다.
    """
    # 경로 보안 검사
    doc_abs = os.path.abspath(doc_path)
    base_abs = os.path.abspath(base_image_dir)
    for p in [doc_abs, base_abs]:
        if not p.lower().startswith(r"c:\ameva"):
            return f"Security Error: Path must be under C:\\ameva. Got: {p}"

    if not os.path.exists(doc_abs):
        return f"Error: Markdown file not found at {doc_path}"
    if not os.path.isdir(base_abs):
        return f"Error: base_image_dir is not a directory: {base_image_dir}"

    # base_image_dir 내 모든 이미지 파일 인덱싱 (파일명 → 절대경로)
    IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".svg", ".webp", ".bmp"}
    image_index = {}
    for root, _, files in os.walk(base_abs):
        for fname in files:
            if any(fname.lower().endswith(ext) for ext in IMAGE_EXTS):
                # 중복 시 첫 번째 발견 우선
                if fname.lower() not in image_index:
                    image_index[fname.lower()] = os.path.join(root, fname)

    with open(doc_abs, "r", encoding="utf-8") as f:
        content = f.read()

    # 마크다운 이미지 패턴: ![alt](path)
    pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
    
    fixed_count = 0
    not_found = []
    
    def replace_path(match):
        nonlocal fixed_count
        alt = match.group(1)
        old_path = match.group(2)
        
        # 이미 유효한 URL 이면 스킵
        if old_path.startswith("http://") or old_path.startswith("https://"):
            return match.group(0)
        
        # 파일명 추출
        img_filename = os.path.basename(old_path).lower()
        
        if img_filename in image_index:
            new_path = image_index[img_filename].replace("\\", "/")
            fixed_count += 1
            return f"![{alt}]({new_path})"
        else:
            not_found.append(old_path)
            return match.group(0)  # 그대로 유지
    
    new_content = pattern.sub(replace_path, content)
    
    if fixed_count == 0 and not not_found:
        return f"No image references found in {doc_path}."
    
    if fixed_count > 0:
        # 수정된 내용 저장
        with open(doc_abs, "w", encoding="utf-8") as f:
            f.write(new_content)
    
    report = (
        f"## 🖼️ MD Image Path Fixer\n\n"
        f"**File**: `{doc_path}`  \n"
        f"**Fixed Paths**: {fixed_count}  \n"
        f"**Not Found**: {len(not_found)}  \n"
        f"**Image Index Size**: {len(image_index)} files indexed\n\n"
    )
    if not_found:
        report += "### ⚠️ Images Not Found (path kept as-is)\n"
        for p in not_found[:10]:
            report += f"- `{p}`\n"
    if fixed_count > 0:
        report += f"\n✅ File saved with {fixed_count} corrected image path(s)."
    
    return report
